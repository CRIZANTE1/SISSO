"""
Gerador de Relatório Word (DOCX) no Padrão Vibra
Replica fielmente o estilo e estrutura do PDF
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
from typing import Dict, Any, List, Optional
import io

# Cor Vibra (cinza usado no PDF)
VIBRA_GRAY = RGBColor(211, 211, 211)  # #d3d3d3
TEXT_GRAY = RGBColor(102, 102, 102)  # #666
LIGHT_GRAY = RGBColor(240, 240, 240)  # #f0f0f0
BORDER_GRAY = RGBColor(204, 204, 204)  # #ccc

def add_vibra_header_cell(cell, text: str):
    """Adiciona célula com estilo vibra-green (fundo cinza, texto em negrito)"""
    p = cell.paragraphs[0]
    p.clear()
    run = p.add_run(text)
    run.font.bold = True
    run.font.size = Pt(9)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'D3D3D3')
    cell._element.get_or_add_tcPr().append(shading)

def add_label_value_cell(cell, label: str, value: str):
    """Adiciona célula com label pequeno e valor"""
    p = cell.paragraphs[0]
    p.clear()
    run_label = p.add_run(f"{label}\n")
    run_label.font.size = Pt(8)
    run_label.font.color.rgb = TEXT_GRAY
    run_label.font.bold = True
    run_value = p.add_run(str(value))
    run_value.font.size = Pt(10)

def add_checkbox_cell(cell, checked: bool, text: str):
    """Adiciona checkbox (☑ ou ☐) com texto"""
    checkbox = "☑" if checked else "☐"
    cell.text = f"{checkbox} {text}"

def generate_word_report(
    accident_data: Dict[str, Any],
    people_data: List[Dict[str, Any]],
    timeline_events: List[Dict[str, Any]],
    verified_causes: List[Dict[str, Any]],
    evidence_images: List[str],
    fault_tree_json: Optional[Dict[str, Any]] = None,
    commission_actions: Optional[List[Dict[str, Any]]] = None,
    image_cache: Optional[Dict[str, str]] = None
) -> bytes:
    """
    Gera relatório em formato Word (.docx) com estilo idêntico ao PDF.
    
    Args:
        accident_data: Dados da tabela 'accidents'
        people_data: Lista da tabela 'involved_people'
        timeline_events: Lista de eventos da timeline
        verified_causes: Lista de nós validados com códigos NBR
        evidence_images: Lista de URLs ou base64 das imagens de evidência
        fault_tree_json: JSON da árvore de falhas (opcional)
        commission_actions: Lista de ações executadas pela comissão (opcional)
        image_cache: Dicionário com cache de imagens pré-carregadas (opcional)
    
    Returns:
        bytes: Documento Word gerado
    """
    try:
        print(f"[WORD_GENERATION] Iniciando geração do Word")
        
        # Cria documento
        doc = Document()
        
        # Configuração de estilo padrão
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(10)
        
        # Prepara data atual
        current_date = datetime.now().strftime('%d/%m/%Y')
        
        # Filtra pessoas por tipo
        commission = [p for p in people_data if p.get('person_type') == 'Commission_Member']
        drivers = [p for p in people_data if p.get('person_type') == 'Driver']
        injured = [p for p in people_data if p.get('person_type') == 'Injured']
        witnesses = [p for p in people_data if p.get('person_type') == 'Witness']
        
        # ========== CAPA ==========
        # Título principal
        title = doc.add_paragraph()
        title_run = title.add_run('RELATÓRIO FINAL')
        title_run.font.size = Pt(28)
        title_run.font.bold = True
        title_run.font.color.rgb = TEXT_GRAY
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.paragraph_format.space_before = Pt(250)
        
        # Subtítulo
        subtitle = doc.add_paragraph()
        subtitle_run = subtitle.add_run('Investigação de Acidente')
        subtitle_run.font.size = Pt(18)
        subtitle_run.font.color.rgb = RGBColor(51, 51, 51)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.paragraph_format.space_before = Pt(30)
        
        # Espaço
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Informações do acidente (centralizadas)
        event_title = accident_data.get('title') or accident_data.get('description', 'N/A')
        info_para = doc.add_paragraph()
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_para.paragraph_format.line_spacing = 2.0
        
        info_para.add_run(f"Evento: {event_title}\n").bold = True
        
        if accident_data.get('site_name'):
            info_para.add_run(f"Base: {accident_data.get('site_name', 'N/A')}\n").bold = True
        
        info_para.add_run(f"Local da Base: {accident_data.get('base_location', 'N/A')}\n").bold = True
        
        occurrence_date = accident_data.get('occurrence_date') or accident_data.get('occurred_at', 'N/A')
        info_para.add_run(f"Data de Ocorrência: {occurrence_date}\n").bold = True
        
        if accident_data.get('created_at'):
            info_para.add_run(f"Comissão Constituída em: {accident_data.get('created_at', 'N/A')}").bold = True
        
        # Quebra de página
        doc.add_page_break()
        
        # ========== METODOLOGIA ==========
        # Título da seção
        section_title = doc.add_paragraph()
        section_title_run = section_title.add_run('2. METODOLOGIA DE INVESTIGAÇÃO')
        section_title_run.font.size = Pt(14)
        section_title_run.font.bold = True
        section_title_run.font.color.rgb = TEXT_GRAY
        section_title.paragraph_format.space_before = Pt(20)
        section_title.paragraph_format.space_after = Pt(10)
        
        # Adiciona linha inferior
        pPr = section_title._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        pBdr_bottom = OxmlElement('w:bottom')
        pBdr_bottom.set(qn('w:val'), 'single')
        pBdr_bottom.set(qn('w:sz'), '16')
        pBdr_bottom.set(qn('w:space'), '1')
        pBdr_bottom.set(qn('w:color'), 'D3D3D3')
        pBdr.append(pBdr_bottom)
        pPr.append(pBdr)
        
        # Metodologia em parágrafos separados (um por seção)
        methodology_paragraphs = [
            "A investigação deste acidente foi conduzida utilizando a metodologia de Análise de Árvore de Falhas (FTA - Fault Tree Analysis), uma técnica sistemática e estruturada para identificar e analisar as causas raiz de eventos indesejados. Esta metodologia permite uma representação gráfica hierárquica das relações causais, facilitando a compreensão dos fatores que contribuíram para a ocorrência do acidente.",
            "A FTA é amplamente reconhecida na indústria como uma ferramenta eficaz para investigação de acidentes, conforme estabelecido pela NBR 14280:2019 - Cadastro de acidente do trabalho - Procedimento e classificação¹, que estabelece os critérios para classificação e registro de acidentes do trabalho no Brasil. A norma define parâmetros para classificação de acidentes com e sem afastamento, além de estabelecer critérios para cálculo de indicadores de segurança.",
            "A metodologia segue os princípios estabelecidos pela ISO 31010:2019 - Risk management - Risk assessment techniques², que apresenta a Análise de Árvore de Falhas como uma técnica recomendada para análise de riscos e investigação de eventos. A norma internacional destaca a capacidade da FTA em identificar combinações de falhas que podem levar a eventos indesejados, permitindo uma análise probabilística e qualitativa dos fatores contribuintes.",
            "Durante a investigação, foram aplicados os princípios da Análise de Causa Raiz (RCA - Root Cause Analysis), conforme metodologia descrita por ABS Consulting (2005)³ em \"Root Cause Analysis Handbook: A Guide to Effective Incident Investigation\". Esta abordagem sistemática permite identificar não apenas as causas imediatas, mas também as causas contribuintes e as causas raiz, garantindo que as ações corretivas sejam direcionadas aos fatores fundamentais que permitiram a ocorrência do evento.",
            "A construção da árvore de falhas seguiu a estrutura hierárquica proposta por Vesely et al. (1981)⁴ em \"Fault Tree Handbook\", onde o evento topo (acidente) é decomposto em eventos intermediários e eventos básicos, utilizando portas lógicas (AND, OR) para representar as relações causais. Esta estrutura permite uma análise sistemática de todas as hipóteses causais, facilitando a validação ou descarte de cada hipótese através de evidências coletadas durante a investigação.",
            "A classificação das causas seguiu os critérios estabelecidos pela NBR 14280:2019¹, que diferencia entre causas imediatas, causas contribuintes e causas básicas. As causas básicas são aquelas que, se corrigidas, podem prevenir a recorrência do acidente, enquanto as causas imediatas são os fatores que diretamente levaram à ocorrência do evento.",
            "A coleta de evidências foi realizada seguindo os princípios de preservação e documentação estabelecidos por OSHA (Occupational Safety and Health Administration)⁵ em suas diretrizes para investigação de acidentes. Todas as evidências foram documentadas, incluindo fotografias, depoimentos, documentos técnicos e registros operacionais, garantindo a rastreabilidade e a confiabilidade das informações utilizadas na análise.",
            "A validação das hipóteses causais foi realizada através da análise crítica das evidências coletadas, seguindo o método científico de formulação e teste de hipóteses. Cada hipótese foi avaliada quanto à sua plausibilidade, consistência com as evidências e capacidade de explicar a sequência de eventos que levou ao acidente, conforme metodologia descrita por Reason (1997)⁶ em \"Managing the Risks of Organizational Accidents\"."
        ]
        
        for text in methodology_paragraphs:
            p = doc.add_paragraph(text)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.line_spacing = 1.6
            p.paragraph_format.space_after = Pt(12)
            p.runs[0].font.size = Pt(10)
        
        # Referências bibliográficas
        doc.add_paragraph()
        ref_title = doc.add_paragraph()
        ref_title_run = ref_title.add_run('REFERÊNCIAS BIBLIOGRÁFICAS')
        ref_title_run.font.bold = True
        ref_title_run.font.size = Pt(9)
        # Aplica fundo cinza
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'D3D3D3')
        ref_title._element.get_or_add_pPr().append(shading)
        
        references = [
            "¹ ABNT - Associação Brasileira de Normas Técnicas. NBR 14280:2019 - Cadastro de acidente do trabalho - Procedimento e classificação. Rio de Janeiro: ABNT, 2019.",
            "² ISO - International Organization for Standardization. ISO 31010:2019 - Risk management - Risk assessment techniques. Geneva: ISO, 2019.",
            "³ ABS Consulting. Root Cause Analysis Handbook: A Guide to Effective Incident Investigation. 3rd ed. Knoxville: ABS Group, 2005.",
            "⁴ VESELY, W. E.; GOLDBERG, F. F.; ROBERTS, N. H.; HAASL, D. F. Fault Tree Handbook. Washington, DC: U.S. Nuclear Regulatory Commission, 1981. (NUREG-0492)",
            "⁵ OSHA - Occupational Safety and Health Administration. Incident Investigation. In: OSHA 2254-09R 2015 - Training Requirements in OSHA Standards and Training Guidelines. Washington, DC: U.S. Department of Labor, 2015.",
            "⁶ REASON, J. Managing the Risks of Organizational Accidents. Aldershot: Ashgate Publishing, 1997."
        ]
        
        for ref in references:
            p = doc.add_paragraph(ref)
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.space_after = Pt(8)
            p.runs[0].font.size = Pt(9)
        
        doc.add_page_break()
        
        # ========== RESUMO GERENCIAL ==========
        # Título da seção
        section_title = doc.add_paragraph()
        section_title_run = section_title.add_run('RESUMO GERENCIAL DO RELATÓRIO FINAL')
        section_title_run.font.size = Pt(14)
        section_title_run.font.bold = True
        section_title_run.font.color.rgb = TEXT_GRAY
        section_title.paragraph_format.space_before = Pt(20)
        section_title.paragraph_format.space_after = Pt(10)
        
        # Adiciona linha inferior
        pPr = section_title._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        pBdr_bottom = OxmlElement('w:bottom')
        pBdr_bottom.set(qn('w:val'), 'single')
        pBdr_bottom.set(qn('w:sz'), '16')
        pBdr_bottom.set(qn('w:space'), '1')
        pBdr_bottom.set(qn('w:color'), 'D3D3D3')
        pBdr.append(pBdr_bottom)
        pPr.append(pBdr)
        
        # Tabela de resumo (estilo form-table)
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        # Primeira linha (cabeçalho não visível, mas primeira linha de dados)
        # Data/Hora
        row = table.rows[0]
        add_vibra_header_cell(row.cells[0], 'Data/Hora')
        row.cells[1].text = str(occurrence_date)
        row.cells[1].paragraphs[0].runs[0].font.size = Pt(10)
        
        # Base
        row = table.add_row()
        add_vibra_header_cell(row.cells[0], 'Base')
        row.cells[1].text = str(accident_data.get('site_name', 'N/A'))
        row.cells[1].paragraphs[0].runs[0].font.size = Pt(10)
        
        # Local da Base
        row = table.add_row()
        add_vibra_header_cell(row.cells[0], 'Local da Base')
        row.cells[1].text = str(accident_data.get('base_location', 'N/A'))
        row.cells[1].paragraphs[0].runs[0].font.size = Pt(10)
        
        # Descrição Resumida
        row = table.add_row()
        add_vibra_header_cell(row.cells[0], 'Descrição Resumida')
        row.cells[1].text = str(accident_data.get('description', accident_data.get('title', 'N/A')))
        row.cells[1].paragraphs[0].runs[0].font.size = Pt(10)
        
        # Tipo
        row = table.add_row()
        add_vibra_header_cell(row.cells[0], 'Tipo')
        row.cells[1].text = str(accident_data.get('type', 'N/A'))
        row.cells[1].paragraphs[0].runs[0].font.size = Pt(10)
        
        # Classificação
        class_parts = []
        if accident_data.get('class_injury'):
            class_parts.append('Com Lesão')
        if accident_data.get('class_community'):
            class_parts.append('Impacto na Comunidade')
        if accident_data.get('class_environment'):
            class_parts.append('Meio Ambiente')
        if accident_data.get('class_process_safety'):
            class_parts.append('Segurança de Processo')
        if accident_data.get('class_asset_damage'):
            class_parts.append('Dano ao Patrimônio')
        if accident_data.get('class_near_miss'):
            class_parts.append('Quase-Acidente')
        
        classification = ', '.join(class_parts) if class_parts else accident_data.get('classification', 'N/A')
        row = table.add_row()
        add_vibra_header_cell(row.cells[0], 'Classificação')
        row.cells[1].text = str(classification)
        row.cells[1].paragraphs[0].runs[0].font.size = Pt(10)
        
        # Ajusta largura das colunas (25% e 75%)
        for row in table.rows:
            row.cells[0].width = Inches(1.5)
            row.cells[1].width = Inches(4.5)
        
        doc.add_page_break()
        
        # ========== INFORMAÇÕES DETALHADAS ==========
        section_title = doc.add_paragraph()
        section_title_run = section_title.add_run('1. INFORMAÇÕES DO EVENTO')
        section_title_run.font.size = Pt(14)
        section_title_run.font.bold = True
        section_title_run.font.color.rgb = TEXT_GRAY
        section_title.paragraph_format.space_before = Pt(20)
        section_title.paragraph_format.space_after = Pt(10)
        
        # Adiciona linha inferior
        pPr = section_title._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        pBdr_bottom = OxmlElement('w:bottom')
        pBdr_bottom.set(qn('w:val'), 'single')
        pBdr_bottom.set(qn('w:sz'), '16')
        pBdr_bottom.set(qn('w:space'), '1')
        pBdr_bottom.set(qn('w:color'), 'D3D3D3')
        pBdr.append(pBdr_bottom)
        pPr.append(pBdr)
        
        # 1.1 Dados Gerais
        sub_title = doc.add_paragraph()
        sub_title_run = sub_title.add_run('1.1. Dados Gerais')
        sub_title_run.font.bold = True
        sub_title_run.font.size = Pt(10)
        # Aplica fundo cinza
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'D3D3D3')
        sub_title._element.get_or_add_pPr().append(shading)
        sub_title.paragraph_format.space_before = Pt(10)
        sub_title.paragraph_format.space_after = Pt(8)
        
        # Tabela de dados gerais (4 colunas)
        general_table = doc.add_table(rows=1, cols=4)
        general_table.style = 'Table Grid'
        
        # Primeira linha
        row = general_table.rows[0]
        add_label_value_cell(row.cells[0], 'Número do Registro', accident_data.get('registry_number', 'N/A'))
        add_label_value_cell(row.cells[1], 'Base', accident_data.get('site_name', 'N/A'))
        add_label_value_cell(row.cells[2], 'Data de Ocorrência', str(occurrence_date))
        add_label_value_cell(row.cells[3], 'Status', accident_data.get('status', 'N/A'))
        
        # Segunda linha (Local da Base - colspan 4)
        row = general_table.add_row()
        cell = row.cells[0]
        # Mescla células
        for i in range(1, 4):
            cell.merge(row.cells[i])
        add_label_value_cell(cell, 'Local da Base', accident_data.get('base_location', 'N/A'))
        
        # Terceira linha (Descrição Completa - colspan 4)
        row = general_table.add_row()
        cell = row.cells[0]
        for i in range(1, 4):
            cell.merge(row.cells[i])
        add_label_value_cell(cell, 'Descrição Completa', accident_data.get('description', accident_data.get('title', 'N/A')))
        
        # 1.2 Classificação
        sub_title = doc.add_paragraph()
        sub_title_run = sub_title.add_run('1.2. Classificação')
        sub_title_run.font.bold = True
        sub_title_run.font.size = Pt(10)
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'D3D3D3')
        sub_title._element.get_or_add_pPr().append(shading)
        sub_title.paragraph_format.space_before = Pt(10)
        sub_title.paragraph_format.space_after = Pt(8)
        
        # Tabela de classificação
        class_table = doc.add_table(rows=1, cols=2)
        class_table.style = 'Table Grid'
        
        row = class_table.rows[0]
        # Coluna 1: Tipo de Impacto
        p1 = row.cells[0].paragraphs[0]
        p1.clear()
        run_label = p1.add_run('Tipo de Impacto\n')
        run_label.font.size = Pt(8)
        run_label.font.color.rgb = TEXT_GRAY
        run_label.font.bold = True
        
        # Checkboxes
        class_injury = accident_data.get('class_injury')
        class_community = accident_data.get('class_community')
        class_environment = accident_data.get('class_environment')
        class_process_safety = accident_data.get('class_process_safety')
        class_asset_damage = accident_data.get('class_asset_damage')
        class_near_miss = accident_data.get('class_near_miss')
        
        checkbox = "☑" if class_injury else "☐"
        p1.add_run(f"{checkbox} Acidente Com Lesão na Força de Trabalho\n")
        checkbox = "☑" if class_community else "☐"
        p1.add_run(f"{checkbox} Acidente Com Lesão na Comunidade\n")
        checkbox = "☑" if class_environment else "☐"
        p1.add_run(f"{checkbox} Impacto ao Meio Ambiente\n")
        checkbox = "☑" if class_process_safety else "☐"
        p1.add_run(f"{checkbox} Segurança de Processo\n")
        checkbox = "☑" if class_asset_damage else "☐"
        p1.add_run(f"{checkbox} Dano ao Patrimônio\n")
        checkbox = "☑" if class_near_miss else "☐"
        p1.add_run(f"{checkbox} Quase-Acidente")
        
        # Coluna 2: Gravidade
        p2 = row.cells[1].paragraphs[0]
        p2.clear()
        run_label = p2.add_run('Gravidade Real/Potencial\n')
        run_label.font.size = Pt(8)
        run_label.font.color.rgb = TEXT_GRAY
        run_label.font.bold = True
        
        severity = accident_data.get('severity_level', '')
        checkbox = "☑" if severity in ['Very Low', 'Muito Baixa'] else "☐"
        p2.add_run(f"{checkbox} Muito Baixa\n")
        checkbox = "☑" if severity in ['Low', 'Baixa'] else "☐"
        p2.add_run(f"{checkbox} Baixa\n")
        checkbox = "☑" if severity in ['Medium', 'Média'] else "☐"
        p2.add_run(f"{checkbox} Média\n")
        checkbox = "☑" if severity in ['High', 'Alta'] else "☐"
        p2.add_run(f"{checkbox} Alta\n")
        checkbox = "☑" if severity in ['Catastrophic', 'Catastrófica'] else "☐"
        p2.add_run(f"{checkbox} Catastrófica")
        
        # Valor estimado de perdas (se houver)
        if accident_data.get('estimated_loss_value'):
            row = class_table.add_row()
            cell = row.cells[0]
            cell.merge(row.cells[1])
            add_label_value_cell(cell, 'Valor Estimado de Perdas', f"R$ {accident_data.get('estimated_loss_value', 0):,.2f}")
        
        # 1.4 Perfil dos Motoristas
        for person in drivers:
            sub_title = doc.add_paragraph()
            sub_title_run = sub_title.add_run(f"1.4. Perfil do Motorista: {person.get('name', 'N/A')}")
            sub_title_run.font.bold = True
            sub_title_run.font.size = Pt(10)
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), 'D3D3D3')
            sub_title._element.get_or_add_pPr().append(shading)
            sub_title.paragraph_format.space_before = Pt(10)
            sub_title.paragraph_format.space_after = Pt(8)
            
            driver_table = doc.add_table(rows=1, cols=5)
            driver_table.style = 'Table Grid'
            
            row = driver_table.rows[0]
            add_label_value_cell(row.cells[0], 'Nome', person.get('name', 'N/A'))
            add_label_value_cell(row.cells[1], 'Idade', str(person.get('age', 'N/A')) if person.get('age') else 'N/A')
            add_label_value_cell(row.cells[2], 'Função', person.get('job_title', 'N/A'))
            add_label_value_cell(row.cells[3], 'Empresa', person.get('company', 'N/A'))
            add_label_value_cell(row.cells[4], 'Matrícula/CPF', person.get('registration_id', 'N/A'))
            
            if person.get('time_in_role') or person.get('aso_date'):
                row = driver_table.add_row()
                cell1 = row.cells[0]
                cell1.merge(row.cells[1])
                add_label_value_cell(cell1, 'Tempo na Função', person.get('time_in_role', 'N/A'))
                cell2 = row.cells[2]
                cell2.merge(row.cells[3])
                cell2.merge(row.cells[4])
                add_label_value_cell(cell2, 'Data ASO', person.get('aso_date', 'N/A'))
        
        # 1.7 Perfil das Vítimas/Lesionados
        for person in injured:
            sub_title = doc.add_paragraph()
            sub_title_run = sub_title.add_run(f"1.7. Perfil da Vítima/Lesionado: {person.get('name', 'N/A')}")
            sub_title_run.font.bold = True
            sub_title_run.font.size = Pt(10)
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), 'D3D3D3')
            sub_title._element.get_or_add_pPr().append(shading)
            sub_title.paragraph_format.space_before = Pt(10)
            sub_title.paragraph_format.space_after = Pt(8)
            
            victim_table = doc.add_table(rows=1, cols=3)
            victim_table.style = 'Table Grid'
            
            # Primeira linha: Nome, Nascimento, Idade
            row = victim_table.rows[0]
            add_label_value_cell(row.cells[0], 'Nome', person.get('name', 'N/A'))
            
            # Formata data de nascimento
            birth_date = person.get('birth_date', '')
            if birth_date:
                try:
                    if isinstance(birth_date, str) and '-' in birth_date:
                        parts = birth_date[:10].split('-')
                        birth_date = f"{parts[2]}/{parts[1]}/{parts[0]}"
                except:
                    pass
            add_label_value_cell(row.cells[1], 'Nascimento', birth_date or 'N/A')
            add_label_value_cell(row.cells[2], 'Idade', str(person.get('age', 'N/A')) if person.get('age') else 'N/A')
            
            # Segunda linha: CPF, RG, Estado Civil
            row = victim_table.add_row()
            add_label_value_cell(row.cells[0], 'CPF', person.get('registration_id', 'N/A'))
            add_label_value_cell(row.cells[1], 'RG', person.get('rg', 'N/A'))
            add_label_value_cell(row.cells[2], 'Estado Civil', person.get('marital_status', 'N/A'))
            
            # Terceira linha: Naturalidade, Número de Filhos
            row = victim_table.add_row()
            add_label_value_cell(row.cells[0], 'Naturalidade', person.get('birthplace', 'N/A'))
            add_label_value_cell(row.cells[1], 'Número de Filhos', str(person.get('children_count', 'N/A')) if person.get('children_count') else 'N/A')
            row.cells[2].text = ''
            
            # Quarta linha: Tipo de Lesão, Parte do Corpo, Dias Perdidos
            row = victim_table.add_row()
            add_label_value_cell(row.cells[0], 'Tipo de Lesão', person.get('injury_type', 'N/A'))
            add_label_value_cell(row.cells[1], 'Parte do Corpo', person.get('body_part', 'N/A'))
            add_label_value_cell(row.cells[2], 'Dias Perdidos', str(person.get('lost_days', 'N/A')) if person.get('lost_days') else 'N/A')
            
            # Quinta linha: Data ASO, N° CAT, Fatalidade
            row = victim_table.add_row()
            add_label_value_cell(row.cells[0], 'Data ASO', person.get('aso_date', 'N/A'))
            add_label_value_cell(row.cells[1], 'N° CAT', person.get('cat_number', 'N/A'))
            p = row.cells[2].paragraphs[0]
            p.clear()
            run_label = p.add_run('Fatalidade\n')
            run_label.font.size = Pt(8)
            run_label.font.color.rgb = TEXT_GRAY
            run_label.font.bold = True
            is_fatal = person.get('is_fatal', False)
            checkbox = "☑" if is_fatal else "☐"
            p.add_run(f"{checkbox} Sim ")
            checkbox = "☑" if not is_fatal else "☐"
            p.add_run(f"{checkbox} Não")
            
            # Sexta linha: Tipo, Função, Empresa
            row = victim_table.add_row()
            p = row.cells[0].paragraphs[0]
            p.clear()
            run_label = p.add_run('Tipo\n')
            run_label.font.size = Pt(8)
            run_label.font.color.rgb = TEXT_GRAY
            run_label.font.bold = True
            emp_type = person.get('employment_type', '')
            checkbox = "☑" if emp_type == 'Empregado' else "☐"
            p.add_run(f"{checkbox} Empregado ")
            checkbox = "☑" if emp_type == 'Contratado' else "☐"
            p.add_run(f"{checkbox} Contratado ")
            checkbox = "☑" if emp_type == 'Terceiros/Comunidade' else "☐"
            p.add_run(f"{checkbox} Terceiros/Comunidade")
            add_label_value_cell(row.cells[1], 'Função', person.get('job_title', 'N/A'))
            add_label_value_cell(row.cells[2], 'Empresa', person.get('company', 'N/A'))
            
            # Sétima linha: Histórico Acidente Anterior
            row = victim_table.add_row()
            cell = row.cells[0]
            for i in range(1, 3):
                cell.merge(row.cells[i])
            p = cell.paragraphs[0]
            p.clear()
            run_label = p.add_run('Histórico Acidente Anterior\n')
            run_label.font.size = Pt(8)
            run_label.font.color.rgb = TEXT_GRAY
            run_label.font.bold = True
            has_history = person.get('previous_accident_history', False)
            checkbox = "☑" if has_history else "☐"
            p.add_run(f"{checkbox} Sim ")
            checkbox = "☑" if not has_history else "☐"
            p.add_run(f"{checkbox} Não")
            
            # Capacitações/Validade (se houver)
            if person.get('certifications'):
                row = victim_table.add_row()
                cell = row.cells[0]
                for i in range(1, 3):
                    cell.merge(row.cells[i])
                p = cell.paragraphs[0]
                p.clear()
                run_label = p.add_run('Capacitações/Validade\n')
                run_label.font.size = Pt(8)
                run_label.font.color.rgb = TEXT_GRAY
                run_label.font.bold = True
                
                certs_str = person.get('certifications', '')
                if ';' in certs_str:
                    for cert_item in certs_str.split(';'):
                        if '|' in cert_item:
                            parts = cert_item.split('|', 1)
                            p.add_run(f"{parts[0]}")
                            p.runs[-1].font.bold = True
                            if parts[1]:
                                validity_dt = parts[1]
                                if isinstance(validity_dt, str) and '-' in validity_dt:
                                    validity_parts = validity_dt[:10].split('-')
                                    validity_dt = f"{validity_parts[2]}/{validity_parts[1]}/{validity_parts[0]}"
                                p.add_run(f" - Validade: {validity_dt}\n")
                        else:
                            p.add_run(f"{cert_item}\n")
                else:
                    p.add_run(certs_str)
            
            # Tempo na Função (se houver)
            if person.get('time_in_role'):
                row = victim_table.add_row()
                cell = row.cells[0]
                for i in range(1, 3):
                    cell.merge(row.cells[i])
                add_label_value_cell(cell, 'Tempo na Função', person.get('time_in_role', 'N/A'))
        
        # 1.8 Testemunhas
        if witnesses:
            sub_title = doc.add_paragraph()
            sub_title_run = sub_title.add_run('1.8. Testemunhas')
            sub_title_run.font.bold = True
            sub_title_run.font.size = Pt(10)
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), 'D3D3D3')
            sub_title._element.get_or_add_pPr().append(shading)
            sub_title.paragraph_format.space_before = Pt(10)
            sub_title.paragraph_format.space_after = Pt(8)
            
            witness_table = doc.add_table(rows=1, cols=3)
            witness_table.style = 'Table Grid'
            
            # Cabeçalho
            header_row = witness_table.rows[0]
            add_vibra_header_cell(header_row.cells[0], 'Nome')
            add_vibra_header_cell(header_row.cells[1], 'Matrícula/CPF')
            add_vibra_header_cell(header_row.cells[2], 'Observações')
            
            for person in witnesses:
                row = witness_table.add_row()
                row.cells[0].text = person.get('name', 'N/A')
                row.cells[0].paragraphs[0].runs[0].font.size = Pt(10)
                row.cells[1].text = person.get('registration_id', 'N/A')
                row.cells[1].paragraphs[0].runs[0].font.size = Pt(10)
                row.cells[2].text = person.get('job_title', 'N/A')
                row.cells[2].paragraphs[0].runs[0].font.size = Pt(10)
        
        # 1.5 Vazamentos / Segurança de Processo
        if accident_data.get('class_process_safety') or accident_data.get('class_environment'):
            sub_title = doc.add_paragraph()
            sub_title_run = sub_title.add_run('1.5. Vazamentos / Segurança de Processo')
            sub_title_run.font.bold = True
            sub_title_run.font.size = Pt(10)
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), 'D3D3D3')
            sub_title._element.get_or_add_pPr().append(shading)
            sub_title.paragraph_format.space_before = Pt(10)
            sub_title.paragraph_format.space_after = Pt(8)
            
            leak_table = doc.add_table(rows=1, cols=4)
            leak_table.style = 'Table Grid'
            
            row = leak_table.rows[0]
            add_label_value_cell(row.cells[0], 'Produto Liberado', accident_data.get('product_released', 'N/A'))
            add_label_value_cell(row.cells[1], 'Volume Liberado', f"{accident_data.get('volume_released', 'N/A')} m³" if accident_data.get('volume_released') else 'N/A')
            add_label_value_cell(row.cells[2], 'Volume Recuperado', f"{accident_data.get('volume_recovered', 'N/A')} m³" if accident_data.get('volume_recovered') else 'N/A')
            add_label_value_cell(row.cells[3], 'Duração', f"{accident_data.get('release_duration_hours', 'N/A')} horas" if accident_data.get('release_duration_hours') else 'N/A')
            
            row = leak_table.add_row()
            cell1 = row.cells[0]
            cell1.merge(row.cells[1])
            add_label_value_cell(cell1, 'Equipamento onde ocorreu a perda de contenção', accident_data.get('equipment_involved', 'N/A'))
            cell2 = row.cells[2]
            cell2.merge(row.cells[3])
            add_label_value_cell(cell2, 'Área Afetada', accident_data.get('area_affected', 'N/A'))
            
            if accident_data.get('process_safety_observation'):
                row = leak_table.add_row()
                cell = row.cells[0]
                for i in range(1, 4):
                    cell.merge(row.cells[i])
                add_label_value_cell(cell, 'Observação', accident_data.get('process_safety_observation', 'N/A'))
        
        # 1.6 Cronologia de Eventos
        if timeline_events:
            sub_title = doc.add_paragraph()
            sub_title_run = sub_title.add_run('1.6. Cronologia de Eventos')
            sub_title_run.font.bold = True
            sub_title_run.font.size = Pt(10)
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), 'D3D3D3')
            sub_title._element.get_or_add_pPr().append(shading)
            sub_title.paragraph_format.space_before = Pt(10)
            sub_title.paragraph_format.space_after = Pt(8)
            
            for event in timeline_events:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.space_after = Pt(10)
                
                time_run = p.add_run(f"{event.get('event_time', 'N/A')}\n")
                time_run.font.bold = True
                time_run.font.color.rgb = TEXT_GRAY
                desc_run = p.add_run(event.get('description', 'N/A'))
                desc_run.font.size = Pt(10)
        
        # 1.7 Cronologia de Ações da Comissão
        if commission_actions:
            sub_title = doc.add_paragraph()
            sub_title_run = sub_title.add_run('1.7. Cronologia de Ações da Comissão')
            sub_title_run.font.bold = True
            sub_title_run.font.size = Pt(10)
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), 'D3D3D3')
            sub_title._element.get_or_add_pPr().append(shading)
            sub_title.paragraph_format.space_before = Pt(15)
            sub_title.paragraph_format.space_after = Pt(8)
            
            info_p = doc.add_paragraph('Abaixo são apresentadas as ações executadas pela comissão durante a investigação.')
            info_p.runs[0].font.size = Pt(9)
            info_p.runs[0].font.color.rgb = TEXT_GRAY
            info_p.paragraph_format.space_after = Pt(10)
            
            for action in commission_actions:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.space_after = Pt(10)
                
                time_run = p.add_run(f"{action.get('action_time', 'N/A')}")
                time_run.font.bold = True
                time_run.font.color.rgb = RGBColor(33, 150, 243)  # #2196f3
                
                if action.get('action_type'):
                    type_run = p.add_run(f" {action.get('action_type')}")
                    type_run.font.bold = True
                    type_run.font.color.rgb = RGBColor(255, 255, 255)
                    # Tenta aplicar fundo azul (limitado no Word)
                
                if action.get('responsible_person'):
                    resp_run = p.add_run(f" Responsável: {action.get('responsible_person')}")
                    resp_run.font.size = Pt(9)
                    resp_run.font.color.rgb = TEXT_GRAY
                
                p.add_run(f"\n{action.get('description', 'N/A')}")
        
        doc.add_page_break()
        
        # ========== ÁRVORE DE FALHAS ==========
        section_title = doc.add_paragraph()
        section_title_run = section_title.add_run('5. ANÁLISE DAS CAUSAS (ÁRVORE DE FALHAS)')
        section_title_run.font.size = Pt(14)
        section_title_run.font.bold = True
        section_title_run.font.color.rgb = TEXT_GRAY
        section_title.paragraph_format.space_before = Pt(20)
        section_title.paragraph_format.space_after = Pt(10)
        
        # Adiciona linha inferior
        pPr = section_title._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        pBdr_bottom = OxmlElement('w:bottom')
        pBdr_bottom.set(qn('w:val'), 'single')
        pBdr_bottom.set(qn('w:sz'), '16')
        pBdr_bottom.set(qn('w:space'), '1')
        pBdr_bottom.set(qn('w:color'), 'D3D3D3')
        pBdr.append(pBdr_bottom)
        pPr.append(pBdr)
        
        doc.add_paragraph('Abaixo a representação gráfica da Árvore de Falhas gerada durante a investigação.')
        
        if fault_tree_json:
            # Adiciona nota sobre a árvore (no Word não podemos renderizar HTML/CSS facilmente)
            note_p = doc.add_paragraph('Nota: A árvore de falhas completa está disponível no sistema e pode ser visualizada na interface web.')
            note_p.runs[0].font.italic = True
            note_p.runs[0].font.color.rgb = TEXT_GRAY
        
        # 5.1 Classificação NBR 14280
        if verified_causes:
            sub_title = doc.add_paragraph()
            sub_title_run = sub_title.add_run('5.1. Classificação NBR 14280')
            sub_title_run.font.size = Pt(14)
            sub_title_run.font.bold = True
            sub_title_run.font.color.rgb = TEXT_GRAY
            sub_title.paragraph_format.space_before = Pt(30)
            sub_title.paragraph_format.space_after = Pt(10)
            
            # Adiciona linha inferior
            pPr = sub_title._element.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            pBdr_bottom = OxmlElement('w:bottom')
            pBdr_bottom.set(qn('w:val'), 'single')
            pBdr_bottom.set(qn('w:sz'), '16')
            pBdr_bottom.set(qn('w:space'), '1')
            pBdr_bottom.set(qn('w:color'), 'D3D3D3')
            pBdr.append(pBdr_bottom)
            pPr.append(pBdr)
            
            nbr_table = doc.add_table(rows=1, cols=3)
            nbr_table.style = 'Table Grid'
            
            # Cabeçalho
            header_row = nbr_table.rows[0]
            add_vibra_header_cell(header_row.cells[0], 'Causa Identificada')
            add_vibra_header_cell(header_row.cells[1], 'Código NBR')
            add_vibra_header_cell(header_row.cells[2], 'Descrição NBR')
            
            for cause in verified_causes:
                row = nbr_table.add_row()
                row.cells[0].text = cause.get('label', 'N/A')
                row.cells[0].paragraphs[0].runs[0].font.size = Pt(9)
                row.cells[1].text = cause.get('nbr_code', 'N/A')
                row.cells[1].paragraphs[0].runs[0].font.bold = True
                row.cells[1].paragraphs[0].runs[0].font.size = Pt(9)
                row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                row.cells[2].text = cause.get('nbr_description', 'N/A')
                row.cells[2].paragraphs[0].runs[0].font.size = Pt(9)
        
        # 5.2 Hipóteses Descritas e Justificadas
        if fault_tree_json:
            from utils.report_generator import extract_hypotheses_from_tree
            hypotheses = extract_hypotheses_from_tree(fault_tree_json)
            
            if hypotheses:
                sub_title = doc.add_paragraph()
                sub_title_run = sub_title.add_run('5.2. Hipóteses Descritas e Justificadas')
                sub_title_run.font.size = Pt(14)
                sub_title_run.font.bold = True
                sub_title_run.font.color.rgb = TEXT_GRAY
                sub_title.paragraph_format.space_before = Pt(30)
                sub_title.paragraph_format.space_after = Pt(10)
                
                # Adiciona linha inferior
                pPr = sub_title._element.get_or_add_pPr()
                pBdr = OxmlElement('w:pBdr')
                pBdr_bottom = OxmlElement('w:bottom')
                pBdr_bottom.set(qn('w:val'), 'single')
                pBdr_bottom.set(qn('w:sz'), '16')
                pBdr_bottom.set(qn('w:space'), '1')
                pBdr_bottom.set(qn('w:color'), 'D3D3D3')
                pBdr.append(pBdr_bottom)
                pPr.append(pBdr)
                
                intro_p = doc.add_paragraph('Abaixo são descritas todas as hipóteses levantadas durante a investigação, com suas respectivas justificativas para descarte ou consideração.')
                intro_p.runs[0].font.size = Pt(10)
                intro_p.paragraph_format.space_after = Pt(15)
                
                for hyp in hypotheses:
                    # Título da hipótese (vibra-green)
                    hyp_title = doc.add_paragraph()
                    node_num = hyp.get('node_number', f'H{hypotheses.index(hyp) + 1}')
                    hyp_label = hyp.get('label', 'N/A')
                    hyp_title_run = hyp_title.add_run(f'Hipótese {node_num}: {hyp_label}')
                    hyp_title_run.font.bold = True
                    hyp_title_run.font.size = Pt(10)
                    shading = OxmlElement('w:shd')
                    shading.set(qn('w:fill'), 'D3D3D3')
                    hyp_title._element.get_or_add_pPr().append(shading)
                    hyp_title.paragraph_format.space_before = Pt(10)
                    hyp_title.paragraph_format.space_after = Pt(8)
                    
                    # Tabela com status, descrição, justificativa
                    hyp_table = doc.add_table(rows=1, cols=2)
                    hyp_table.style = 'Table Grid'
                    
                    # Status
                    row = hyp_table.rows[0]
                    row.cells[0].text = 'Status'
                    row.cells[0].paragraphs[0].runs[0].font.bold = True
                    row.cells[0].paragraphs[0].runs[0].font.size = Pt(9)
                    shading = OxmlElement('w:shd')
                    shading.set(qn('w:fill'), 'F0F0F0')
                    row.cells[0]._element.get_or_add_tcPr().append(shading)
                    
                    status = hyp.get('status', 'pending')
                    status_text = ''
                    if status == 'discarded':
                        status_text = '✕ DESCARTADA'
                    elif status == 'validated':
                        status_text = '✓ CONSIDERADA/VALIDADA'
                    else:
                        status_text = '⏳ EM ANÁLISE'
                    
                    p_status = row.cells[1].paragraphs[0]
                    p_status.clear()
                    status_run = p_status.add_run(status_text)
                    status_run.font.bold = True
                    if status == 'discarded':
                        status_run.font.color.rgb = RGBColor(211, 47, 47)  # #d32f2f
                    elif status == 'validated':
                        status_run.font.color.rgb = RGBColor(46, 125, 50)  # #2e7d32
                    else:
                        status_run.font.color.rgb = TEXT_GRAY
                    status_run.font.size = Pt(9)
                    
                    # Descrição
                    row = hyp_table.add_row()
                    row.cells[0].text = 'Descrição'
                    row.cells[0].paragraphs[0].runs[0].font.bold = True
                    row.cells[0].paragraphs[0].runs[0].font.size = Pt(9)
                    shading = OxmlElement('w:shd')
                    shading.set(qn('w:fill'), 'F0F0F0')
                    row.cells[0]._element.get_or_add_tcPr().append(shading)
                    row.cells[1].text = hyp_label
                    row.cells[1].paragraphs[0].runs[0].font.size = Pt(9)
                    
                    # Justificativa
                    row = hyp_table.add_row()
                    row.cells[0].text = 'Justificativa'
                    row.cells[0].paragraphs[0].runs[0].font.bold = True
                    row.cells[0].paragraphs[0].runs[0].font.size = Pt(9)
                    shading = OxmlElement('w:shd')
                    shading.set(qn('w:fill'), 'F0F0F0')
                    row.cells[0]._element.get_or_add_tcPr().append(shading)
                    
                    justification = hyp.get('justification', '')
                    if justification:
                        row.cells[1].text = justification
                        row.cells[1].paragraphs[0].runs[0].font.size = Pt(9)
                    elif status == 'discarded':
                        row.cells[1].text = 'Hipótese descartada após análise da evidência disponível.'
                        row.cells[1].paragraphs[0].runs[0].font.italic = True
                        row.cells[1].paragraphs[0].runs[0].font.color.rgb = TEXT_GRAY
                        row.cells[1].paragraphs[0].runs[0].font.size = Pt(9)
                    elif status == 'validated':
                        row.cells[1].text = 'Hipótese validada com base nas evidências coletadas durante a investigação.'
                        row.cells[1].paragraphs[0].runs[0].font.italic = True
                        row.cells[1].paragraphs[0].runs[0].font.color.rgb = TEXT_GRAY
                        row.cells[1].paragraphs[0].runs[0].font.size = Pt(9)
                    else:
                        row.cells[1].text = 'Hipótese em análise, aguardando validação ou descarte.'
                        row.cells[1].paragraphs[0].runs[0].font.italic = True
                        row.cells[1].paragraphs[0].runs[0].font.color.rgb = TEXT_GRAY
                        row.cells[1].paragraphs[0].runs[0].font.size = Pt(9)
                    
                    # Código NBR (se houver)
                    if hyp.get('nbr_code'):
                        row = hyp_table.add_row()
                        row.cells[0].text = 'Código NBR'
                        row.cells[0].paragraphs[0].runs[0].font.bold = True
                        row.cells[0].paragraphs[0].runs[0].font.size = Pt(9)
                        shading = OxmlElement('w:shd')
                        shading.set(qn('w:fill'), 'F0F0F0')
                        row.cells[0]._element.get_or_add_tcPr().append(shading)
                        nbr_text = f"{hyp.get('nbr_code', 'N/A')} - {hyp.get('nbr_description', '')}"
                        p_nbr = row.cells[1].paragraphs[0]
                        p_nbr.clear()
                        code_run = p_nbr.add_run(hyp.get('nbr_code', 'N/A'))
                        code_run.font.bold = True
                        p_nbr.add_run(f" - {hyp.get('nbr_description', '')}")
                        row.cells[1].paragraphs[0].runs[0].font.size = Pt(9)
                    
                    doc.add_paragraph()  # Espaço entre hipóteses
        
        doc.add_page_break()
        
        # ========== RECOMENDAÇÕES ==========
        if fault_tree_json:
            from utils.report_generator import extract_recommendations_from_tree
            recommendations = extract_recommendations_from_tree(fault_tree_json)
            
            if recommendations.get('basic_causes') or recommendations.get('contributing_causes'):
                section_title = doc.add_paragraph()
                section_title_run = section_title.add_run('9. RECOMENDAÇÕES')
                section_title_run.font.size = Pt(14)
                section_title_run.font.bold = True
                section_title_run.font.color.rgb = TEXT_GRAY
                section_title.paragraph_format.space_before = Pt(20)
                section_title.paragraph_format.space_after = Pt(10)
                
                # Adiciona linha inferior
                pPr = section_title._element.get_or_add_pPr()
                pBdr = OxmlElement('w:pBdr')
                pBdr_bottom = OxmlElement('w:bottom')
                pBdr_bottom.set(qn('w:val'), 'single')
                pBdr_bottom.set(qn('w:sz'), '16')
                pBdr_bottom.set(qn('w:space'), '1')
                pBdr_bottom.set(qn('w:color'), 'D3D3D3')
                pBdr.append(pBdr_bottom)
                pPr.append(pBdr)
                
                intro_p = doc.add_paragraph('Abaixo são apresentadas as recomendações para prevenir ou corrigir as causas básicas e contribuintes identificadas na investigação.')
                intro_p.runs[0].font.size = Pt(10)
                intro_p.paragraph_format.space_after = Pt(15)
                
                if recommendations.get('basic_causes'):
                    sub_title = doc.add_paragraph()
                    sub_title_run = sub_title.add_run('9.1. Recomendações para Causas Básicas')
                    sub_title_run.font.bold = True
                    sub_title_run.font.size = Pt(10)
                    shading = OxmlElement('w:shd')
                    shading.set(qn('w:fill'), 'D3D3D3')
                    sub_title._element.get_or_add_pPr().append(shading)
                    sub_title.paragraph_format.space_before = Pt(15)
                    sub_title.paragraph_format.space_after = Pt(10)
                    
                    for i, rec in enumerate(recommendations.get('basic_causes', []), 1):
                        # Título da recomendação
                        rec_title = doc.add_paragraph()
                        rec_title_run = rec_title.add_run(f"{i}. {rec.get('label', 'N/A')}")
                        rec_title_run.font.bold = True
                        rec_title_run.font.size = Pt(11)
                        rec_title_run.font.color.rgb = TEXT_GRAY
                        rec_title.paragraph_format.space_before = Pt(0)
                        rec_title.paragraph_format.space_after = Pt(5)
                        
                        # Código NBR (se houver)
                        if rec.get('nbr_code'):
                            nbr_p = doc.add_paragraph()
                            nbr_run = nbr_p.add_run(f"Código NBR: {rec.get('nbr_code', 'N/A')} - {rec.get('nbr_description', '')}")
                            nbr_run.font.size = Pt(9)
                            nbr_run.font.color.rgb = TEXT_GRAY
                            nbr_p.paragraph_format.space_after = Pt(8)
                        
                        # Texto da recomendação (com borda esquerda e fundo)
                        rec_p = doc.add_paragraph(rec.get('recommendation', 'Nenhuma recomendação fornecida.'))
                        rec_p.paragraph_format.left_indent = Inches(0.5)
                        rec_p.paragraph_format.space_before = Pt(5)
                        rec_p.paragraph_format.space_after = Pt(6)
                        # Aplica fundo cinza claro
                        shading = OxmlElement('w:shd')
                        shading.set(qn('w:fill'), 'F9F9F9')
                        rec_p._element.get_or_add_pPr().append(shading)
                        # Adiciona borda esquerda (limitado no Word, mas tentamos)
                        pBdr = OxmlElement('w:pBdr')
                        pBdr_left = OxmlElement('w:left')
                        pBdr_left.set(qn('w:val'), 'single')
                        pBdr_left.set(qn('w:sz'), '32')  # 4px
                        pBdr_left.set(qn('w:space'), '1')
                        pBdr_left.set(qn('w:color'), 'D3D3D3')
                        pBdr.append(pBdr_left)
                        rec_p._element.get_or_add_pPr().append(pBdr)
                        rec_p.runs[0].font.size = Pt(10)
                        
                        doc.add_paragraph()  # Espaço entre recomendações
                
                if recommendations.get('contributing_causes'):
                    sub_title = doc.add_paragraph()
                    sub_title_run = sub_title.add_run('9.2. Recomendações para Causas Contribuintes')
                    sub_title_run.font.bold = True
                    sub_title_run.font.size = Pt(10)
                    shading = OxmlElement('w:shd')
                    shading.set(qn('w:fill'), 'D3D3D3')
                    sub_title._element.get_or_add_pPr().append(shading)
                    sub_title.paragraph_format.space_before = Pt(20)
                    sub_title.paragraph_format.space_after = Pt(10)
                    
                    for i, rec in enumerate(recommendations.get('contributing_causes', []), 1):
                        # Título da recomendação
                        rec_title = doc.add_paragraph()
                        rec_title_run = rec_title.add_run(f"{i}. {rec.get('label', 'N/A')}")
                        rec_title_run.font.bold = True
                        rec_title_run.font.size = Pt(11)
                        rec_title_run.font.color.rgb = TEXT_GRAY
                        rec_title.paragraph_format.space_before = Pt(0)
                        rec_title.paragraph_format.space_after = Pt(5)
                        
                        # Código NBR (se houver)
                        if rec.get('nbr_code'):
                            nbr_p = doc.add_paragraph()
                            nbr_run = nbr_p.add_run(f"Código NBR: {rec.get('nbr_code', 'N/A')} - {rec.get('nbr_description', '')}")
                            nbr_run.font.size = Pt(9)
                            nbr_run.font.color.rgb = TEXT_GRAY
                            nbr_p.paragraph_format.space_after = Pt(8)
                        
                        # Texto da recomendação (com borda esquerda azul e fundo)
                        rec_p = doc.add_paragraph(rec.get('recommendation', 'Nenhuma recomendação fornecida.'))
                        rec_p.paragraph_format.left_indent = Inches(0.5)
                        rec_p.paragraph_format.space_before = Pt(5)
                        rec_p.paragraph_format.space_after = Pt(6)
                        # Aplica fundo cinza claro
                        shading = OxmlElement('w:shd')
                        shading.set(qn('w:fill'), 'F9F9F9')
                        rec_p._element.get_or_add_pPr().append(shading)
                        # Adiciona borda esquerda azul
                        pBdr = OxmlElement('w:pBdr')
                        pBdr_left = OxmlElement('w:left')
                        pBdr_left.set(qn('w:val'), 'single')
                        pBdr_left.set(qn('w:sz'), '32')  # 4px
                        pBdr_left.set(qn('w:space'), '1')
                        pBdr_left.set(qn('w:color'), '2196F3')  # Azul
                        pBdr.append(pBdr_left)
                        rec_p._element.get_or_add_pPr().append(pBdr)
                        rec_p.runs[0].font.size = Pt(10)
                        
                        doc.add_paragraph()  # Espaço entre recomendações
        
        # ========== COMISSÃO ==========
        if commission:
            doc.add_page_break()
            section_title = doc.add_paragraph()
            section_title_run = section_title.add_run('7. COMISSÃO DE INVESTIGAÇÃO')
            section_title_run.font.size = Pt(14)
            section_title_run.font.bold = True
            section_title_run.font.color.rgb = TEXT_GRAY
            section_title.paragraph_format.space_before = Pt(20)
            section_title.paragraph_format.space_after = Pt(10)
            
            # Adiciona linha inferior
            pPr = section_title._element.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            pBdr_bottom = OxmlElement('w:bottom')
            pBdr_bottom.set(qn('w:val'), 'single')
            pBdr_bottom.set(qn('w:sz'), '16')
            pBdr_bottom.set(qn('w:space'), '1')
            pBdr_bottom.set(qn('w:color'), 'D3D3D3')
            pBdr.append(pBdr_bottom)
            pPr.append(pBdr)
            
            commission_table = doc.add_table(rows=1, cols=4)
            commission_table.style = 'Table Grid'
            
            # Cabeçalho
            header_row = commission_table.rows[0]
            add_vibra_header_cell(header_row.cells[0], 'Nome')
            add_vibra_header_cell(header_row.cells[1], 'Cargo/Função')
            add_vibra_header_cell(header_row.cells[2], 'Matrícula/ID')
            add_vibra_header_cell(header_row.cells[3], 'Participação')
            
            for member in commission:
                row = commission_table.add_row()
                row.cells[0].text = member.get('name', 'N/A')
                row.cells[0].paragraphs[0].runs[0].font.size = Pt(9)
                row.cells[1].text = member.get('job_title', 'N/A')
                row.cells[1].paragraphs[0].runs[0].font.size = Pt(9)
                row.cells[2].text = member.get('registration_id', 'N/A')
                row.cells[2].paragraphs[0].runs[0].font.size = Pt(9)
                row.cells[3].text = member.get('commission_role') or member.get('training_status') or 'Membro da Comissão'
                row.cells[3].paragraphs[0].runs[0].font.size = Pt(9)
        
        # Salva em bytes
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        
        print(f"[WORD_GENERATION] ✓ Word gerado com sucesso: {len(doc_bytes.getvalue())} bytes")
        return doc_bytes.getvalue()
        
    except Exception as e:
        print(f"[WORD_GENERATION] ✗ Erro ao gerar Word: {str(e)}")
        import traceback
        traceback.print_exc()
        raise Exception(f"Erro ao gerar Word: {str(e)}")
